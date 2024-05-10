#!/usr/bin/env python
# -*- coding: utf-8 -*-
# @Time         : 2024/5/6 15:46
# @Author       : adolf
# @Email        : adolf1321794021@gmail.com
# @LastEditTime : 2024/5/6 15:46
# @File         : 生成.py
import io
import json

import streamlit as st
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from loguru import logger

from src.utils import main_chain
from unstructure_my.tools.doc2json import docx_to_json

_page_config = {"label": "生成", "icon": "🌟"}

st.set_page_config(page_title="生成", page_icon="🌟", layout="wide")

st.title("合同智能生成系统")


def load_json_from_file(file_path):
    """从指定的文件路径加载JSON数据"""
    with open(file_path, "r", encoding="utf-8") as file:
        return json.load(file)


# st.write('This is page 2.')
def add_heading(document, text, level):
    # 检查样式是否已存在
    style_name = f"Heading{level}"
    if style_name not in document.styles:
        # 如果样式不存在，则添加
        style = document.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = document.styles["Normal"]
        font = style.font
        if level == 1:
            font.size = Pt(18)  # 或其他你想要的字号
            # 设置其他属性如bold等...
        elif level == 2:
            font.size = Pt(16)
            # 同上，设置其他属性
        else:
            font.size = Pt(14)

    # 使用已存在的样式
    paragraph = document.add_paragraph(text, style=style_name)
    return paragraph


def add_paragraph(document, text):
    document.add_paragraph(text)


def add_list_item(document, text):
    document.add_paragraph(text, style="List Bullet")


def generate_document_from_json(
    json_data, ck_data, AName, ANumber, BName, BNumber, Where
):
    document = Document()

    # 添加合同标题
    add_heading(document, json_data["title"], 1)

    # 添加合同双方基本信息
    add_paragraph(document, f"采购方(以下简称甲方):")
    add_paragraph(document, f"供应方(以下简称乙方):")

    # logger.info(json_data)
    # 遍历子节点，添加相应内容
    for node in json_data["sub_nodes"]:
        if "type" not in node:
            # logger.info(node)
            add_heading(document, node["title"], 2)
            add_paragraph(document, node["content"])
        elif node["type"] == "Node":
            add_heading(document, node["title"], 2)
            if "sub_nodes" in node:
                for sub_node in node["sub_nodes"]:
                    add_heading(document, sub_node["title"], 3)
                    if "type" not in sub_node:
                        add_paragraph(document, sub_node["content"])

                    elif sub_node["type"] == "KeyPoint":
                        if "prompt" not in node:
                            continue
                        for key, value in ck_data.items():
                            if sub_node["prompt"] in key:
                                input_text = (
                                    f"请你根据如下内容{value},回答问题{sub_node['paramsPrompt']}"
                                )
                                response = main_chain().invoke(
                                    {"input": input_text},
                                    # {"callbacks": [st_callback]},
                                )
                                logger.info(response)

                                # content = sub_node['template'].format(**sub_node['params'])
                                content = ""
                                add_paragraph(document, content)

                    elif sub_node["type"] == "Retrieve":
                        for key, value in ck_data.items():
                            if sub_node["prompt"] in key:
                                add_paragraph(document, value)
                    elif sub_node["type"] == "FromSt":
                        logger.info(sub_node["template"])
                        logger.info(AName, ANumber, BName, BNumber, Where)
                        content = (
                            f"{sub_node['template']}".replace("ANAME", AName)
                            .replace("ANUMBER", ANumber)
                            .replace("BNAME", BName)
                            .replace("BNUMBER", BNumber)
                            .replace("WHERE", Where)
                        )
                        # content = fsub_node["template"].format(AName, ANumber, BName, BNumber, Where)
                        add_paragraph(document, content)

                    else:
                        logger.info(sub_node)
                        # add_paragraph(document, sub_node["content"])

        elif node["type"] == "KeyPoint":
            add_heading(document, node["title"], level=2)
            if "prompt" not in node:
                continue
            for key, value in ck_data.items():
                if node["prompt"] in key:
                    input_text = (
                        f"请你根据如下内容{value},回答问题{node['paramsPrompt']}，直接给出答案，不用有多余的内容"
                    )
                    response = main_chain().invoke(
                        {"input": input_text},
                        # {"callbacks": [st_callback]},
                    )
                    logger.info(response)
                    # content = sub_node['template'].format(**sub_node['params'])
                    content = response["text"] + node["template"]
                    add_paragraph(document, content)
            add_paragraph(document, node["template"])
        elif node["type"] == "Retrieve":
            add_heading(document, node["title"], level=2)
            for key, value in ck_data.items():
                if node["prompt"] in key:
                    add_paragraph(document, value)

        else:  # 其他类型节点的处理
            logger.info(node)
            # add_paragraph(document, node['content'])
            content = node.get("content", "")
            if content:
                add_paragraph(document, content)

    # 保存文档
    # document.save(file_path)
    # docx_bytes = docx.Document.save(document, "output.docx")
    # print(f"Document generated at {file_path}")
    with io.BytesIO() as buffer:
        document.save(buffer)
        docx_bytes = buffer.getvalue()

    return docx_bytes


def generate_document_from_file(
    json_file_path, ck_data, AName, ANumber, BName, BNumber, Where
):
    """从JSON文件生成Word文档"""
    json_data = load_json_from_file(json_file_path)
    docx_bytes = generate_document_from_json(
        json_data, ck_data, AName, ANumber, BName, BNumber, Where
    )
    return docx_bytes


uploaded_file = st.file_uploader("请选择一个docx文件进行上传", type=["docx"])

if uploaded_file is not None:
    data = docx_to_json(uploaded_file)
    print(data)

    AName = st.text_input("甲方负责人是谁?", key="AName")
    ANumber = st.text_input("甲方联系方式是什么?", key="ANumber")
    BName = st.text_input("乙方负责人是谁？", key="BName")
    BNumber = st.text_input("乙方联系方式是什么？", key="BNumber")
    Where = st.text_input("实施地点是在哪？", key="Where")
    logger.info(AName, ANumber, BName, BNumber, Where)

    if AName and ANumber and BName and BNumber and Where:
        # 指定JSON文件路径和输出Word文档路径
        json_file_path = "web_show/procure_01.json"
        # output_docx_path = "download.docx"

        # 从JSON文件生成Word文档
        document = generate_document_from_file(
            json_file_path, data, AName, ANumber, BName, BNumber, Where
        )
        # print(f"Document generated from {json_file_path} and saved at {output_docx_path}")
        st.download_button(
            label="获取生成的文件",
            data=document,
            file_name="generated_document.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
